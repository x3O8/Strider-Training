from __future__ import annotations

import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


def run_meta(run) -> str:
    flags: list[str] = []
    if run.bold:
        flags.append("bold")
    if run.italic:
        flags.append("italic")
    if run.underline:
        flags.append("underline")
    if run.font.strike:
        flags.append("strike")
    if run.font.color and run.font.color.rgb:
        flags.append(f"color={run.font.color.rgb}")
    if run.font.highlight_color:
        flags.append(f"highlight={run.font.highlight_color}")
    return ",".join(flags) or "plain"


def paragraph_text(paragraph) -> str:
    parts = []
    for run in paragraph.runs:
        if run.text:
            parts.append(f"<{run_meta(run)}>{run.text}</>")
    return "".join(parts) or paragraph.text


def extract(path: Path) -> str:
    doc = Document(path)
    output = [f"DOCUMENT: {path.name}", "", "PARAGRAPHS:"]
    for index, paragraph in enumerate(doc.paragraphs, 1):
        if paragraph.text.strip() or paragraph.runs:
            output.append(
                f"P{index:03d} [{paragraph.style.name if paragraph.style else 'None'}] "
                f"{paragraph_text(paragraph)}"
            )

    if doc.tables:
        output.extend(["", "TABLES:"])
        for table_index, table in enumerate(doc.tables, 1):
            output.append(f"TABLE {table_index}")
            for row_index, row in enumerate(table.rows, 1):
                cells = []
                for cell in row.cells:
                    cells.append(" / ".join(p.text for p in cell.paragraphs if p.text.strip()))
                output.append(f"R{row_index:03d}: " + " | ".join(cells))

    with zipfile.ZipFile(path) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))
        output.extend(["", "RAW XML PARAGRAPHS (includes tracked insertions/deletions):"])
        for index, paragraph in enumerate(document_xml.iter(W + "p"), 1):
            parts = []
            for node in paragraph.iter():
                if node.tag == W + "t" and node.text:
                    parts.append(node.text)
                elif node.tag == W + "delText" and node.text:
                    parts.append(f"[DELETED:{node.text}]")
                elif node.tag == W + "tab":
                    parts.append("\t")
                elif node.tag == W + "br":
                    parts.append("\n")
            text = "".join(parts).strip()
            if text:
                output.append(f"X{index:03d}: {text}")

        if "word/comments.xml" in archive.namelist():
            comments_xml = ET.fromstring(archive.read("word/comments.xml"))
            output.extend(["", "COMMENTS:"])
            for comment in comments_xml.findall(W + "comment"):
                comment_text = "".join(
                    node.text or "" for node in comment.iter(W + "t")
                ).strip()
                output.append(
                    f"COMMENT {comment.get(W + 'id')} by {comment.get(W + 'author')}: {comment_text}"
                )

    return "\n".join(output)


if __name__ == "__main__":
    source = Path(sys.argv[1])
    print(extract(source))
